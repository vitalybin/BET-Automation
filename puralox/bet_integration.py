# puralox/bet_integration.py
from __future__ import annotations
from pathlib import Path
from typing import Dict, List, Tuple, Union
from datetime import datetime
import re
import json

import fitz  # PyMuPDF
import pandas as pd
import matplotlib.pyplot as plt

try:
    from docx import Document
    from docx.shared import Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    HAVE_DOCX = True
except Exception:
    Document = None
    Inches = None
    WD_ALIGN_PARAGRAPH = None
    HAVE_DOCX = False


def _open_pdf(pdf_source: Union[str, Path, bytes]) -> fitz.Document:
    if isinstance(pdf_source, (bytes, bytearray)):
        return fitz.open(stream=pdf_source, filetype="pdf")
    p = Path(pdf_source).expanduser()
    if not p.exists():
        raise FileNotFoundError(f"PDF not found: {p}")
    return fitz.open(str(p))

# ---------- Helpers ----------
def _unique_preserve(seq: List[str]) -> List[str]:
    seen = set()
    out = []
    for s in seq:
        key = (s or "").strip().casefold()
        if key and key not in seen:
            seen.add(key)
            out.append((s or "").strip())
    return out

def _norm_spaces(s: str | None) -> str:
    if not s:
        return ""
    return re.sub(r"\s+", " ", str(s)).strip()

# ---------- Highlight helpers (unchanged API) ----------
def _iter_quad_rects(annot) -> List[fitz.Rect]:
    rects: List[fitz.Rect] = []
    verts = getattr(annot, "vertices", None) or getattr(annot, "quad_points", None)
    if verts:
        for i in range(0, len(verts), 4):
            quad = verts[i:i+4]
            xs, ys = [], []
            for pt in quad:
                if hasattr(pt, "x") and hasattr(pt, "y"):
                    xs.append(float(pt.x)); ys.append(float(pt.y))
                elif isinstance(pt, (tuple, list)) and len(pt) >= 2:
                    xs.append(float(pt[0])); ys.append(float(pt[1]))
            if xs and ys:
                rects.append(fitz.Rect(min(xs), min(ys), max(xs), max(ys)))
    if not rects:
        try:
            rects = [annot.rect]
        except Exception:
            rects = []
    return rects

def extract_highlight_strings(pdf_source: Union[str, Path, bytes]) -> List[str]:
    doc = _open_pdf(pdf_source)
    highlights: List[Tuple[int, float, float, str]] = []
    for page_index in range(len(doc)):
        page = doc[page_index]
        words = page.get_text("words") or []
        word_entries = [{
            "rect": fitz.Rect(w[0], w[1], w[2], w[3]),
            "text": w[4],
            "x": float(w[0]),
            "y": float(w[1]),
        } for w in words]

        annot = page.first_annot
        while annot:
            subtype = None
            try:
                subtype = annot.type[0]
            except Exception:
                pass
            if subtype == fitz.PDF_ANNOT_HIGHLIGHT:
                quad_rects = _iter_quad_rects(annot)
                captured = [
                    we for we in word_entries
                    if any(we["rect"].intersects(qr) for qr in quad_rects)
                ]
                captured.sort(key=lambda w: (round(w["y"], 1), w["x"]))
                text = " ".join(w["text"] for w in captured).strip()
                if text:
                    y = captured[0]["y"] if captured else 0.0
                    x = captured[0]["x"] if captured else 0.0
                    highlights.append((page_index, y, x, text))
            annot = annot.next

    highlights.sort(key=lambda t: (t[0], t[1], t[2]))
    return [t[3] for t in highlights]

FIELD_PATTERNS: List[Tuple[str, List[re.Pattern]]] = [
    ("Operator", [re.compile(r"\bOperator:\s*([^\n]+)", re.I)]),
    ("Date", [re.compile(r"\bDate:\s*([0-9/\-:. ]+)", re.I)]),
    ("Sample ID", [re.compile(r"\bSample ID:\s*([^\n]+)", re.I)]),
    ("Filename", [re.compile(r"\bFilename:\s*([^\n]+)", re.I)]),
    ("Sample Description", [re.compile(r"\bSample Desc:\s*([^\n]+)", re.I)]),
    ("Comment", [re.compile(r"\bComment:\s*([^\n]+)", re.I)]),
    ("Sample weight", [re.compile(r"\bSample weight:\s*([^\n]+)", re.I)]),
    ("Sample Volume", [re.compile(r"\bSample Volume:\s*([^\n]+)", re.I)]),
    ("Outgas Time", [re.compile(r"\bOutgas Time:\s*([^\n]+)", re.I)]),
    ("OutgasTemp", [re.compile(r"\bOutgasTemp:\s*([^\n]+)", re.I)]),
    ("Analysis gas", [re.compile(r"\bAnalysis gas:\s*([^\n]+)", re.I)]),
    ("Bath Temp", [re.compile(r"\bBath Temp:\s*([^\n]+)", re.I)]),
    ("Press. Tolerance", [re.compile(r"\bPress\. Tolerance:\s*([^\n]+)", re.I)]),
    ("Equil time", [re.compile(r"\bEquil time:\s*([^\n]+)", re.I)]),
    ("Equil timeout", [re.compile(r"\bEquil timeout:\s*([^\n]+)", re.I)]),
    ("Analysis Time", [re.compile(r"\bAnalysis Time:\s*([^\n]+)", re.I)]),
    ("End of run", [re.compile(r"\bEnd of run:\s*([^\n]+)", re.I)]),
    ("Instrument (verbatim)", [re.compile(r"\bInstrument:\s*([^\n]+)", re.I)]),
    ("Cell ID", [re.compile(r"\bCell ID:\s*([^\n]+)", re.I)]),
    ("Isotherm Slope", [re.compile(r"\bSlope\s*=\s*([^\s]+)", re.I)]),
    ("Isotherm Intercept", [re.compile(r"\bIntercept\s*=\s*([^\s]+)", re.I)]),
    ("Isotherm r", [re.compile(r"\bCorrelation coefficient,\s*r\s*=\s*([^\s]+)", re.I)]),
    ("C constant", [re.compile(r"\bC constant\s*=\s*([^\s]+)", re.I)]),
    ("Surface Area", [re.compile(r"\bSurface Area\s*=\s*([^\n]+)", re.I)]),
    ("Thickness method", [re.compile(r"\bThickness method:\s*([^\n]+)", re.I)]),
    ("Multi-Point BET Slope", [re.compile(r"\bMulti-Point BET.*?Slope\s*=\s*([^\s]+)", re.I)]),
    ("Multi-Point BET Intercept", [re.compile(r"\bMulti-Point BET.*?Intercept\s*=\s*([^\s]+)", re.I)]),
    ("Multi-Point BET r", [re.compile(r"\bMulti-Point BET.*?r\s*=\s*([^\s]+)", re.I)]),
    ("Micropore volume", [re.compile(r"\bMicropore volume\s*=\s*([^\n]+)", re.I)]),
    ("Micropore area", [re.compile(r"\bMicropore area\s*=\s*([^\n]+)", re.I)]),
    ("External surface area", [re.compile(r"\bExternal surface area\s*=\s*([^\n]+)", re.I)]),
    ("t-Plot Surface Area", [re.compile(r"\bSurface Area\s*=\s*([^\n]+)", re.I)]),
    ("Pore Volume", [re.compile(r"\bPore Volume\s*=\s*([^\n]+)", re.I)]),
    ("Pore Diameter Dv(d)", [re.compile(r"\bPore Diameter Dv\(d\)\s*=\s*([^\n]+)", re.I)]),
    ("BJH Avg Pore Diameter", [re.compile(r"\bAverage Pore Diameter\s*[:=]\s*([^\n]+)", re.I)]),
    ("BJH Pore Volume", [re.compile(r"\bBJH Pore Volume\s*[:=]\s*([^\n]+)", re.I)]),
    ("BJH Pore Surface Area", [re.compile(r"\bPore Surface Area\s*[:=]\s*([^\n]+)", re.I)]),
]

def _expand_snippet(snippet: str) -> List[str]:
    h = re.sub(r"\s+", " ", snippet).strip()
    parts = re.split(
        r"(?<=\S)\s{2,}(?=\S)|(?= Operator:)|(?= Sample )|(?= Analysis )|"
        r"(?= Surface Area)|(?= Slope =)|(?= Intercept =)|(?= C constant=)|"
        r"(?= Correlation coefficient)|(?= Thickness method:)",
        h
    )
    return [p.strip() for p in parts if p.strip()]

def parse_highlights(highlights: List[str]) -> Dict[str, str]:
    data: Dict[str, str] = {}
    expanded: List[str] = []
    for h in highlights:
        expanded.extend(_expand_snippet(h))

    for snippet in expanded:
        matched = False
        for key, patterns in FIELD_PATTERNS:
            if key in data and len(data[key]) > 2:
                continue
            for pat in patterns:
                m = pat.search(snippet)
                if m:
                    data.setdefault(key, _norm_spaces(m.group(1)))
                    matched = True
                    break
            if matched:
                break

        if not matched and ":" in snippet:
            k, v = [s.strip() for s in snippet.split(":", 1)]
            if k and v and k not in data:
                data[k] = _norm_spaces(v)

    for k, v in list(data.items()):
        data[k] = _norm_spaces(v)
    return data

# ---------- Full-text parsers ----------
def _full_text(doc: fitz.Document) -> str:
    return "\n".join(doc[i].get_text() for i in range(len(doc)))

def _section(text: str, start_pat: str, end_pat: str | None, flags=re.I | re.S) -> str:
    m1 = re.search(start_pat, text, flags)
    m2 = re.search(end_pat, text, flags) if end_pat else None
    if not m1:
        return ""
    start = m1.end()
    end = m2.start() if m2 else len(text)
    if m2 and end < start:
        end = len(text)
    return text[start:end]

def parse_general(text: str) -> Dict[str, object]:
    d: Dict[str, object] = {}
    ops = [ _norm_spaces(o) for o in re.findall(r"\bOperator:\s*([^\n]+)", text, re.I) ]
    ops = _unique_preserve([o for o in ops if o])
    d["Operators"] = ops
    d["OperatorPrimary"] = ops[0] if ops else ""

    d["Dates"] = re.findall(r"\bDate[: ]\s*([0-9]{4}[/\-][0-9]{2}[/\-][0-9]{2}(?:\s[0-9:]+)?)", text, re.I)

    def grab(label: str, out_key: str | None = None):
        key = out_key or label
        m = re.search(rf"\b{re.escape(label)}:\s*([^\n]*)", text, re.I)
        if m:
            d[key] = _norm_spaces(m.group(1))

    for lab in ["Sample ID","Filename","Sample Desc","Comment","Sample weight","Sample Volume","Outgas Time","OutgasTemp",
                "Analysis gas","Bath Temp","Press. Tolerance","Equil time","Equil timeout","Analysis Time","End of run","Instrument","Cell ID"]:
        grab(lab, "Sample Description" if lab=="Sample Desc" else None)
    return d

def parse_isotherm_summary(text: str) -> Dict[str, str]:
    blk = _section(text, r"\bIsotherm\b", r"\bMBET summary\b|Page 2 of") or text
    out: Dict[str, str] = {}
    for key, pat in [
        ("Isotherm Slope", r"Slope\s*=\s*([^\s]+)"),
        ("Isotherm Intercept", r"Intercept\s*=\s*([^\s]+)"),
        ("Isotherm r", r"Correlation coefficient,\s*r\s*=\s*([^\s]+)"),
        ("C constant", r"C constant\s*=\s*([^\s]+)"),
        ("Surface Area", r"Surface Area\s*=\s*([0-9.+\-eE ]+(?:m²|m2)/g)"),
    ]:
        m = re.search(pat, blk, re.I)
        if m:
            out[key] = _norm_spaces(m.group(1))
    return out

def parse_isotherm_points(text: str) -> pd.DataFrame | None:
    region = _section(text, r"Relative Pressure, P/Po", r"Isotherm\s*[\r\n ]*Slope") \
          or _section(text, r"Relative Pressure, P/Po", r"MBET summary|Page 2 of")
    floats = re.findall(r"[-+]?\d+(?:\.\d+)?(?:e[+-]?\d+)?", region or "", re.I)
    vals = [float(x) for x in floats]
    pairs: List[Tuple[float, float]] = []
    for i in range(0, len(vals) - 1, 2):
        ppo, vol = vals[i], vals[i + 1]
        if 0.0 <= ppo <= 1.2 and 5 <= vol <= 2000:
            pairs.append((ppo, vol))
    if not pairs:
        return None
    df = pd.DataFrame(pairs, columns=["P_over_P0", "Vol_cc_g_STP"])
    df = df.sort_values(["P_over_P0", "Vol_cc_g_STP"]).drop_duplicates("P_over_P0", keep="last")
    df = df.sort_values("P_over_P0").reset_index(drop=True)
    return df

def parse_multipoint_bet_summary(text: str) -> Dict[str, str]:
    blk = _section(text, r"Multi-Point BET", r"t plot|t-plot|t plot") or text
    out: Dict[str, str] = {}
    for key, pat in [
        ("Thickness method", r"Thickness method:\s*([^\n]+)"),
        ("Multi-Point BET Slope", r"Slope\s*=\s*([^\s]+)"),
        ("Multi-Point BET Intercept", r"Intercept\s*=\s*([^\s]+)"),
        ("Multi-Point BET r", r"Correlation coefficient,\s*r\s*=\s*([^\s]+)"),
        ("Micropore volume", r"Micropore volume\s*=\s*([^\n]+)"),
        ("Micropore area", r"Micropore area\s*=\s*([^\n]+)"),
        ("External surface area", r"External surface area\s*=\s*([^\n]+)"),
    ]:
        m = re.search(pat, blk, re.I)
        if m: out[key] = _norm_spaces(m.group(1))
    return out

def parse_multipoint_bet_table(text: str) -> pd.DataFrame | None:
    region = _section(text, r"Multi-Point BET Plot", r"Multi-Point BET|t plot|t-plot")
    if not region:
        return None
    vals = [float(x) for x in re.findall(r"[-+]?\d+(?:\.\d+)?(?:e[+-]?\d+)?", region, re.I)]
    rows: List[Tuple[float, float, float]] = []
    for i in range(0, len(vals) - 2, 3):
        a, b, c = vals[i:i+3]
        if 0.05 <= a <= 0.8 and 20 <= b <= 1500 and 0.05 <= c <= 10:
            rows.append((a, b, c))
    if len(rows) < 3:
        return None
    df = pd.DataFrame(rows, columns=["P_over_P0", "Vol_cc_g_STP", "BET_transform"])
    df = df.sort_values("P_over_P0").drop_duplicates("P_over_P0", keep="last").reset_index(drop=True)
    return df

def parse_tplot_summary_and_points(text: str) -> tuple[Dict[str, str], pd.DataFrame | None]:
    blk = _section(text, r"\bt plot\b|\bt-plot\b", r"BJH|BJH desorption") or text
    out: Dict[str, str] = {}
    for key, pat in [
        ("t-Plot Surface Area", r"Surface Area\s*=\s*([^\n]+)"),
        ("Pore Volume", r"Pore Volume\s*=\s*([^\n]+)"),
        ("Pore Diameter Dv(d)", r"Pore Diameter Dv\(d\)\s*=\s*([^\n]+)"),
    ]:
        m = re.search(pat, blk, re.I)
        if m: out[key] = _norm_spaces(m.group(1))

    region = _section(text, r"Statistical Thickness \(nm\)", r"BJH|BJH desorption") or blk
    vals = [float(x) for x in re.findall(r"[-+]?\d+(?:\.\d+)?(?:e[+-]?\d+)?", region, re.I)]
    rows: List[Tuple[float, float]] = []
    for i in range(0, len(vals) - 2, 3):
        ppo, tnm, vol = vals[i:i+3]
        if 0.05 <= ppo <= 0.7 and 0.2 <= tnm <= 1.5 and 50 <= vol <= 1500:
            rows.append((tnm, vol))
    if len(rows) < 3:
        return out, None
    df = pd.DataFrame(rows, columns=["Thickness_nm", "Volume_cc_g_STP"]).sort_values("Thickness_nm").reset_index(drop=True)
    return out, df

def parse_bjh_table(text: str) -> pd.DataFrame | None:
    lines = text.splitlines()
    start = None
    for i, ln in enumerate(lines):
        if "Pore Diameter (nm)" in ln:
            start = i
            break
    if start is None:
        return None
    unit_re = re.compile(r"^\s*\[.*\]\s*$")
    unit_lines = [i for i in range(start, min(len(lines), start + 200)) if unit_re.match(lines[i].strip())]
    scan_start = (unit_lines[-1] + 1) if unit_lines else (start + 1)

    num_re = re.compile(r"^[-+]?\d+(?:\.\d+)?(?:e[+-]?\d+)?\s*$", re.I)
    nums: List[float] = []
    for i in range(scan_start, min(len(lines), scan_start + 800)):
        s = lines[i].strip()
        if not s:
            continue
        if num_re.match(s):
            try:
                nums.append(float(s))
            except Exception:
                pass
        else:
            if len(nums) >= 7:
                break

    rows = []
    for i in range(0, len(nums), 7):
        chunk = nums[i:i+7]
        if len(chunk) == 7:
            rows.append(chunk)
    if not rows:
        return None

    df = pd.DataFrame(rows, columns=[
        "Diameter_nm", "PoreVol_ccg", "PoreArea_m2g", "dV_d", "dS_d", "dV_logd", "dS_logd"
    ]).sort_values("Diameter_nm").reset_index(drop=True)
    return df

# ---------- Plots ----------
def plot_isotherm(df: pd.DataFrame, out_png: Path) -> None:
    plt.figure()
    plt.plot(df["P_over_P0"], df["Vol_cc_g_STP"], marker=".")
    plt.xlabel("Relative Pressure (P/P0)")
    plt.ylabel("Relative Volume @ STP (cc/g)")
    plt.title("Isotherm")
    plt.savefig(out_png, bbox_inches="tight", dpi=160)
    plt.close()

def plot_tplot(df: pd.DataFrame, out_png: Path) -> None:
    plt.figure()
    plt.plot(df["Thickness_nm"], df["Volume_cc_g_STP"], marker="s")
    plt.xlabel("Statistical Thickness (nm)")
    plt.ylabel("Volume @ STP (cc/g)")
    plt.title("t-Plot: Volume vs Thickness")
    plt.savefig(out_png, bbox_inches="tight", dpi=160)
    plt.close()

def plot_bjh(df: pd.DataFrame, out_png: Path) -> None:
    plt.figure()
    plt.plot(df["Diameter_nm"], df["dV_logd"], marker="o")
    plt.xlabel("Pore Diameter (nm)")
    plt.ylabel("dV(log d) (cc/g)")
    plt.title("BJH Desorption: dV(log d) vs Diameter")
    plt.savefig(out_png, bbox_inches="tight", dpi=160)
    plt.close()

def build_docx(bundle: Dict[str, object], dfs: Dict[str, pd.DataFrame], plots: List[Path], out_docx: Path) -> None:
    if not HAVE_DOCX:
        return
    doc = Document()
    doc.add_heading("BET Analysis Report (Auto-Extracted)", level=0)
    p = doc.add_paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for sec_key, title in [
        ("general", "General"),
        ("isotherm_summary", "BET Isotherm Summary"),
        ("multipoint_bet_summary", "Multi-Point BET Summary"),
        ("tplot_summary", "t-Plot Summary"),
    ]:
        sec = bundle.get(sec_key, {})
        if sec:
            doc.add_heading(title, level=1)
            for k, v in sec.items():
                doc.add_paragraph(f"{k}: {v}")
    if plots:
        doc.add_heading("Plots", level=1)
        for pp in plots:
            doc.add_picture(str(pp), width=Inches(6))
    doc.save(out_docx)

# ---------- High-level ----------
def extract_all_with_prints(pdf_source: Union[str, Path, bytes],
                            out_dir: Union[str, Path] = "bet_out",
                            skip_docx: bool = True) -> Dict[str, object]:
    out_dir = Path(out_dir); out_dir.mkdir(parents=True, exist_ok=True)

    doc = _open_pdf(pdf_source)
    text = _full_text(doc)

    general = parse_general(text)
    iso_summary = parse_isotherm_summary(text)
    iso_points = parse_isotherm_points(text)
    mp_summary = parse_multipoint_bet_summary(text)
    mp_table = parse_multipoint_bet_table(text)
    tp_summary, tp_points = parse_tplot_summary_and_points(text)
    bjh_table = parse_bjh_table(text)

    csvs = {}
    if iso_points is not None:
        p = out_dir / "isotherm_points.csv"; iso_points.to_csv(p, index=False); csvs["isotherm"] = str(p)
    if mp_table is not None:
        p = out_dir / "multipoint_bet.csv"; mp_table.to_csv(p, index=False); csvs["multipoint_bet"] = str(p)
    if tp_points is not None:
        p = out_dir / "tplot_points.csv"; tp_points.to_csv(p, index=False); csvs["tplot"] = str(p)
    if bjh_table is not None:
        p = out_dir / "bjh_desorption.csv"; bjh_table.to_csv(p, index=False); csvs["bjh"] = str(p)

    plots: List[Path] = []
    if iso_points is not None and not iso_points.empty:
        p = out_dir / "isotherm.png"; plot_isotherm(iso_points, p); plots.append(p)
    if tp_points is not None and not tp_points.empty:
        p = out_dir / "tplot.png"; plot_tplot(tp_points, p); plots.append(p)
    if bjh_table is not None and not bjh_table.empty:
        p = out_dir / "bjh_distribution.png"; plot_bjh(bjh_table, p); plots.append(p)

    if (not skip_docx) and HAVE_DOCX:
        docx_path = out_dir / "BET_Full_Extract_Report.docx"
        dfs_preview = {
            "Isotherm points": iso_points,
            "Multi-Point BET table": mp_table,
            "t-Plot points": tp_points,
            "BJH Desorption table": bjh_table,
        }
        build_docx(
            {
                "general": general,
                "isotherm_summary": iso_summary,
                "multipoint_bet_summary": mp_summary,
                "tplot_summary": tp_summary,
            },
            dfs_preview, plots, docx_path
        )
        docx_out = str(docx_path)
    else:
        docx_out = "Skipped"

    bundle = {
        "general": general,
        "isotherm_summary": iso_summary,
        "multipoint_bet_summary": mp_summary,
        "tplot_summary": tp_summary,
        "csvs": csvs,
        "plots": [str(p) for p in plots],
        "docx": docx_out,
    }
    (out_dir / "bet_full_extract.json").write_text(json.dumps(bundle, indent=2), encoding="utf-8")
    return bundle
